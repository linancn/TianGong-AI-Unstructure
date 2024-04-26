from weaviate import Client
import json
import weaviate
from weaviate import AdditionalConfig
from docx import Document


## Read the object by ID-V3 method ##
# client = Client("http://localhost:8080")
# data_object = client.data_object.get_by_id(
#     "007dc5fe-1fa1-4013-bddb-3dc2235cf8bc",
#     class_name="Water",
# )

# data_object_json = json.dumps(data_object, ensure_ascii=False, indent=2)
# print(data_object_json)

## Read properties and ids,get the amount of objects in the collection-V3 method ##


client = Client("http://localhost:8080")


def get_batch_with_cursor(collection_name, batch_size, cursor=None):
    try:
        query = (
            client.query.get(
                collection_name,
                ["title", "content", "source"],  # update with the required properties
            )
            .with_additional(["id"])
            .with_limit(batch_size)
        )

        if cursor is not None:
            result = query.with_after(cursor).do()
        else:
            result = query.do()

        return result["data"]["Get"][collection_name]
    except Exception as e:
        print(f"Error occurred at cursor: {cursor}")
        print(e)
        return []


collection_names = ["Water_pptx"]
total_results_fetched = 0

doc = Document()

for collection_name in collection_names:
    cursor = None
    while True:
        next_batch = get_batch_with_cursor(collection_name, 100, cursor)
        batch_count = len(next_batch)
        total_results_fetched += batch_count
        if len(next_batch) == 0:
            break
        for obj in next_batch:
            doc.add_paragraph(str(obj))
            doc.add_paragraph("-" * 50)

        # Move the cursor to the last returned uuid
        cursor = next_batch[-1]["_additional"]["id"]

print(f"Total number of results fetched: {total_results_fetched}")

doc.save("output.docx")
## Delete all the objects in the weaviate  ##
# w_client = weaviate.connect_to_local(
#     host="localhost", additional_config=AdditionalConfig(timeout=(60000, 80000))
# )
# w_client.collections.delete(name="Water_pptx")
# w_client.close()
# ## Vector similarity search with text ##
# import weaviate.classes as wvc

# jeopardy = w_client.collections.get("JeopardyQuestion")
# response = jeopardy.query.near_text(
#     query="animals in movies",
#     limit=2,
#     return_metadata=wvc.query.MetadataQuery(distance=True)
# )

# for o in response.objects:
#         print(o.properties)
#         print(o.metadata.distance)
