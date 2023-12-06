from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
from dotenv import load_dotenv
from openai import OpenAI
from unstructured.chunking.title import chunk_by_title
from unstructured.documents.elements import CompositeElement, Table
from unstructured.partition.docx import partition_docx

from tools.func_calling import get_formatted_text

load_dotenv()

client = OpenAI()


def generate_docx(text_list):
    keys_list = [list(d.keys())[0] for d in text_list]
    value_list = [list(d.values())[0] for d in text_list]

    doc = Document()

    # 设置无间距样式
    no_spacing_style = doc.styles["No Spacing"]
    font = no_spacing_style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # 设置中文字体
    no_spacing_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 设置段落格式
    paragraph_format = no_spacing_style.paragraph_format
    paragraph_format.line_spacing = 1.15  # 行间距
    paragraph_format.space_before = Pt(0)  # 段前间距
    paragraph_format.space_after = Pt(0)  # 段后间距

    # 创建二级标题和无间距段落
    for key, value in zip(keys_list, value_list):
        # 创建二级标题
        doc.add_heading(key, level=2)

        # 创建无间距段落
        formatted_value = get_formatted_text(value)
        paragraph = doc.add_paragraph(formatted_value)
        paragraph.style = no_spacing_style
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 两端对齐

    # 保存文档
    doc.save("output.docx")


def extract_text(file_name: str):
    elements = partition_docx(
        filename=file_name,
        multipage_sections=True,
        infer_table_structure=True,
        include_page_breaks=False,
    )

    chunks = chunk_by_title(
        elements=elements,
        multipage_sections=True,
        combine_text_under_n_chars=0,
        new_after_n_chars=None,
        max_characters=4096,
    )

    text_list = []

    for chunk in chunks:
        if isinstance(chunk, CompositeElement):
            text = chunk.text
            text_list.append(text)
        elif isinstance(chunk, Table):
            if text_list:
                text_list[-1] = text_list[-1] + "\n" + chunk.text
            else:
                text_list.append(chunk.text)
    result_list = []
    for text in text_list:
        split_text = text.split("\n\n", 1)
        if len(split_text) == 2:
            title, content = split_text
        result_list.append({title: content})
    return result_list


file_name = "raw/谢经良. 污水处理设备操作维护问答. 化学工业出版社, 2012.docx"
contents = extract_text(file_name)
generate_docx(contents)
