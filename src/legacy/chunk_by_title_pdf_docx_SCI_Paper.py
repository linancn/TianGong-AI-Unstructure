import tempfile
from unstructured.chunking.title import chunk_by_title
from unstructured.cleaners.core import (
    clean,
    group_broken_paragraphs,
)
from unstructured.documents.elements import Footer, Header, Image, CompositeElement, Table
from unstructured.partition.auto import partition
from tools.vision import vision_completion
from docx import Document

# PDF文件名
pdf_name = "raw/4.pdf"

# 图像的最小尺寸要求
min_image_width = 250
min_image_height = 270

# 关键词列表，用于识别需要过滤的部分
keywords_for_misc = [
    "References",
    "REFERENCES",
    "Bibliography",
    "BIBLIOGRAPHY",
    "Acknowledgements",
    "ACKNOWLEDGEMENTS",
    "Acknowledgments",
    "ACKNOWLEDGMENTS",
    "参考文献",
    "致谢",
    "謝辞",
    "謝",
    "Online",
    "Declaration of Competing Interest",
    "Data availability",
    "Acknowledgments",
    "Supplementary materials",
]

# 分割文档
elements = partition(
    filename=pdf_name,
    header_footer=False,
    pdf_extract_images=False,
    pdf_image_output_dir_path=tempfile.gettempdir(),
    skip_infer_table_types=["jpg", "png", "xls", "xlsx"],
    strategy="auto",
)

# 过滤掉页眉和页脚
filtered_elements = [
    element
    for element in elements
    if not (isinstance(element, Header) or isinstance(element, Footer))
]

# 对文本和图像元素进行处理
for element in filtered_elements:
    if element.text != "":
        element.text = group_broken_paragraphs(element.text)
        element.text = clean(
            element.text,
            bullets=False,
            extra_whitespace=True,
            dashes=False,
            trailing_punctuation=False,
        )
    elif isinstance(element, Image):
        point1 = element.metadata.coordinates.points[0]
        point2 = element.metadata.coordinates.points[2]
        width = abs(point2[0] - point1[0])
        height = abs(point2[1] - point1[1])
        if width >= min_image_width and height >= min_image_height:
            element.text = vision_completion(element.metadata.image_path)

# 将文档分割成块
chunks = chunk_by_title(
    elements=filtered_elements,
    multipage_sections=True,
    combine_text_under_n_chars=0,
    new_after_n_chars=None,
    max_characters=4096,
)

# 寻找包含关键词的标题的块的索引
index_to_remove = None
for i, chunk in enumerate(chunks):
    if isinstance(chunk, CompositeElement):
        # 提取标题，假设是文本的第一行
        title = chunk.text.split('\n')[0] if chunk.text else ""
        if any(keyword in title for keyword in keywords_for_misc):
            index_to_remove = i
            break

# 如果找到包含关键词的标题，则删除该块及之后的所有块
if index_to_remove is not None:
    chunks = chunks[:index_to_remove]

# 打印剩余块的内容
for chunk in chunks:
    if isinstance(chunk, CompositeElement):
        title = chunk.text.split('\n')[0] if chunk.text else ""
        body = '\n'.join(chunk.text.split('\n')[1:]) if chunk.text else ""
        print(f"Title: {title}\nBody: {body}\n")
    else:
        print(chunk.text)
    print("\n" + "-" * 80 + "\n")
# 创建一个新的 Word 文档
doc = Document()

chunks = chunks[:index_to_remove]

# 打印剩余块的内容并将其添加到 Word 文档中
for chunk in chunks:
    if isinstance(chunk, CompositeElement):
        title = chunk.text.split('\n')[0] if chunk.text else ""
        body = '\n'.join(chunk.text.split('\n')[1:]) if chunk.text else ""
        print(f"Title: {title}\nBody: {body}\n")
        doc.add_heading(title, level=1)
        doc.add_paragraph(body)
    else:
        print(chunk.text)
        doc.add_paragraph(chunk.text)
    print("\n" + "-" * 80 + "\n")

# 保存 Word 文档
doc.save("output.docx")